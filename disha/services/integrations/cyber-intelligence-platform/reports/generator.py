"""
Report generator — produces a .docx Cyber Intelligence Report.

Requires: pip install python-docx
"""
import logging
import os
from datetime import datetime, timezone

logger = logging.getLogger(__name__)


def generate_report(data: list, output_dir: str = ".") -> str:
    """
    Generate a Word document report from a list of alert dicts.

    Args:
        data: List of alert dicts (as returned by alert_engine.generate_alert).
        output_dir: Directory to save the report. Defaults to CWD.

    Returns:
        Absolute path to the saved report file.
    """
    try:
        from docx import Document  # python-docx
    except ImportError:
        logger.error("python-docx is not installed. Run: pip install python-docx")
        raise

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    filename = f"cyber_intelligence_report_{timestamp}.docx"
    output_path = os.path.abspath(os.path.join(output_dir, filename))

    doc = Document()
    doc.add_heading("🚨 Disha Cyber Intelligence Report", 0)
    doc.add_paragraph(f"Generated: {timestamp}")
    doc.add_paragraph(f"Total alerts: {len(data)}")
    doc.add_heading("Alerts", level=1)

    for i, item in enumerate(data, start=1):
        doc.add_heading(f"Alert #{i}", level=2)
        if isinstance(item, dict):
            for key, value in item.items():
                doc.add_paragraph(f"{key}: {value}")
        else:
            doc.add_paragraph(str(item))

    doc.save(output_path)
    logger.info("Report saved to: %s", output_path)
    return output_path
