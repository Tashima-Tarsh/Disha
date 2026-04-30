from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT_PATH = Path(__file__).resolve().parent / "DISHA_JARVIS_X_TDD.docx"
ACCENT = RGBColor(79, 70, 229)
ACCENT_2 = RGBColor(56, 189, 248)
TEXT = RGBColor(31, 41, 55)
MUTED = RGBColor(100, 116, 139)
LIGHT_BG = "EEF2FF"
CODE_BG = "F8FAFC"
RULE_BG = "E0F2FE"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_margins(section) -> None:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT

    for style_name, size in [
        ("Title", 28),
        ("Heading 1", 18),
        ("Heading 2", 13),
        ("Heading 3", 11),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.color.rgb = ACCENT
        style.font.bold = True
        style.font.size = Pt(size)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DISHA + JARVIS-X")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(30)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Technical Design Document")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = TEXT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Unified reference architecture for the DISHA repository and the JARVIS-X personal AI, "
        "security, and monitoring platform."
    )
    r.font.size = Pt(11.5)
    r.font.color.rgb = MUTED

    doc.add_paragraph("")
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("Document Type", "Enterprise Technical Design Document"),
        ("Scope", "DISHA repository architecture + JARVIS-X target platform"),
        (
            "Primary Themes",
            "AI reasoning, secure execution, memory, monitoring, deception, deployment",
        ),
        ("Audience", "Architects, engineers, security leads, operators"),
        ("Status", "Design baseline grounded in current repository state"),
    ]
    for row, (left, right) in zip(table.rows, rows):
        row.cells[0].text = left
        row.cells[1].text = right
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(row.cells[0], LIGHT_BG)

    doc.add_page_break()


def add_section_intro(doc: Document, title: str, text: str) -> None:
    doc.add_heading(title, level=1)
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_kv_table(
    doc: Document, rows: list[tuple[str, str]], title: str | None = None
) -> None:
    if title:
        doc.add_paragraph(title).runs[0].bold = True
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Area"
    hdr[1].text = "Details"
    for cell in hdr:
        set_cell_shading(cell, LIGHT_BG)
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_code_block(doc: Document, title: str, code: str) -> None:
    doc.add_paragraph(title).runs[0].bold = True
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CODE_BG)
    p = cell.paragraphs[0]
    r = p.add_run(code.strip())
    r.font.name = "Consolas"
    r.font.size = Pt(8.8)
    r.font.color.rgb = TEXT


def add_rule_panel(doc: Document, rules: list[str]) -> None:
    doc.add_paragraph("No Failure Architecture Rules").runs[0].bold = True
    table = doc.add_table(rows=len(rules), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row, rule in zip(table.rows, rules):
        row.cells[0].text = rule
        set_cell_shading(row.cells[0], RULE_BG)


def add_footer(section, label: str) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(label)
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED


def build() -> None:
    doc = Document()
    set_page_margins(doc.sections[0])
    style_document(doc)
    add_footer(doc.sections[0], "DISHA + JARVIS-X TDD")

    add_title_page(doc)

    add_section_intro(
        doc,
        "1. Executive Summary",
        "This document unifies the architecture of the DISHA repository and the JARVIS-X platform into a single "
        "technical design reference. DISHA is treated as a single platform under consolidation. JARVIS-X is treated as "
        "the target orchestration architecture that brings reasoning, execution, memory, security, monitoring, "
        "and deception-aware telemetry into one coherent system.",
    )
    add_bullets(
        doc,
        [
            "DISHA contributes the current web hardening path, CLI runtime, AI core, and legacy service modules.",
            "JARVIS-X contributes the bounded multi-brain architecture and the additive product workspace under `disha/`.",
            "The combined design is privacy-first, zero-trust, event-driven, and realistic for local or cloud deployment.",
        ],
    )

    add_section_intro(
        doc,
        "2. Repository Scope",
        "The repository contains multiple runtime families. The technical baseline must acknowledge them explicitly rather "
        "than pretending they are already one fully converged platform.",
    )
    add_kv_table(
        doc,
        [
            (
                "web/",
                "Next.js control plane with auth, RBAC, CSRF, rate limiting, file workflows, audit, and shares",
            ),
            (
                "src/",
                "TypeScript CLI and MCP execution gateway with secure storage policy and audit hooks",
            ),
            (
                "backend/",
                "Legacy FastAPI backend with agents, multimodal flows, ranking, RL, and streaming",
            ),
            (
                "disha/ai/core",
                "Reasoning, memory, cognitive loop, citations, and decision engine substrate",
            ),
            (
                "disha-agi-brain/",
                "Prototype AI platform backend and model-routing support",
            ),
            (
                "disha/",
                "Production-minded additive workspace for backend brains, agent, mobile scaffold, and docs",
            ),
        ],
        title="Primary repository surfaces",
    )

    add_section_intro(
        doc,
        "3. Target Unified Architecture",
        "JARVIS-X is the north-star architecture for the repo. The repo should converge around a secure experience layer, "
        "an orchestration layer, a cognitive brain layer, a security and deception layer, a telemetry layer, a data layer, "
        "and an edge sensor layer.",
    )
    add_kv_table(
        doc,
        [
            (
                "Experience Layer",
                "Web console, CLI, mobile companion, voice entry, alerts",
            ),
            (
                "Orchestration Layer",
                "API gateway, request router, session coordinator, event workflow manager",
            ),
            (
                "Cognitive Brain Layer",
                "Reasoning, planning, execution, memory, intelligence, and policy-aware decisioning",
            ),
            (
                "Security and Deception Layer",
                "Token auth, RBAC, safe execution, anomaly detection, risk, decisions, honeypots",
            ),
            (
                "Telemetry Layer",
                "Host metrics, process metrics, network signals, model outputs, deception event enrichment",
            ),
            (
                "Data and Memory Layer",
                "SQLite or Postgres, Redis, audit events, memory, telemetry and risk logs",
            ),
            (
                "Edge and Sensor Layer",
                "Desktop agent, future mobile agent, honeypot collectors, health probes",
            ),
        ],
    )

    add_section_intro(
        doc,
        "4. Multi-Brain System Design",
        "The brain system is modular but not disconnected. Each brain has a bounded role, explicit inputs and outputs, "
        "and a place inside the layered platform.",
    )
    add_kv_table(
        doc,
        [
            (
                "Reasoning Brain",
                "Interprets intent, normalizes goals, and produces structured task meaning",
            ),
            ("Planner", "Converts reasoning output into ordered executable steps"),
            (
                "Execution Brain",
                "Runs bounded tools with confirmation and workspace controls",
            ),
            (
                "Memory Brain",
                "Stores short-term session context and long-term user and system memory",
            ),
            (
                "Security Brain",
                "Applies allow, ask, block, monitor, limit, and isolate decisions",
            ),
            (
                "Intelligence Brain",
                "Retrieves knowledge, selects models, validates outputs, and supports explanation",
            ),
            (
                "Deception Brain",
                "Sub-function of the Security Brain that handles honeypot signals and deception telemetry",
            ),
        ],
    )

    add_section_intro(
        doc,
        "5. Repo-to-Architecture Mapping",
        "The target system reuses current modules instead of discarding them. This reduces delivery risk and preserves existing investment.",
    )
    add_kv_table(
        doc,
        [
            (
                "Reasoning + Memory Kernel",
                "`disha/ai/core/cognitive_loop.py`, `disha/ai/core/memory/*`, `disha/ai/core/intelligence/*`",
            ),
            (
                "Secure Web Control Plane",
                "`web/lib/server/*`, `web/services/*`, `web/app/api/*`",
            ),
            (
                "Execution Gateway",
                "`src/entrypoints/mcp.ts`, `src/security/*`, `src/observability/*`",
            ),
            (
                "Legacy Specialist Services",
                "`backend/app/*` and `disha-agi-brain/backend/*` for staged selective reuse",
            ),
            (
                "JARVIS-X Product Workspace",
                "`disha/brain/*`, `disha/edge_agent/*`, `disha/mobile/*`",
            ),
        ],
    )

    add_section_intro(
        doc,
        "6. Security, Honeypots, and Defensive Deception",
        "The architecture includes deception as a first-class defensive signal source. Honeypots do not perform offensive action. "
        "They amplify attacker behavior signals and enrich the Security Brain’s risk context.",
    )
    add_bullets(
        doc,
        [
            "Honeypots must be isolated from real production assets.",
            "Canary tokens and synthetic credentials must never overlap with real credential paths.",
            "Deception telemetry feeds anomaly correlation and alert enrichment.",
            "High-confidence signals may drive MONITOR, LIMIT, ASK, or ISOLATE recommendations, but not destructive auto-remediation.",
        ],
    )

    add_section_intro(
        doc,
        "7. Data Model and Memory Strategy",
        "The combined platform uses durable and ephemeral storage in distinct roles. Short-term context belongs in session memory. "
        "Long-term preferences, events, and risk posture belong in persistent storage.",
    )
    add_kv_table(
        doc,
        [
            ("Short-Term Memory", "In-request context and Redis-backed session state"),
            (
                "Long-Term Memory",
                "Postgres in DISHA web path and SQLite in JARVIS-X workspace baseline",
            ),
            ("Event Log", "Command, telemetry, and system events"),
            ("Risk Log", "Risk level, score, action, reasons, correlation ids"),
            (
                "Telemetry Store",
                "CPU, memory, process count, send/receive volumes, active app",
            ),
        ],
    )
    add_code_block(
        doc,
        "Representative SQLite schema excerpt",
        """
create table memory (...);
create table events (...);
create table risk_logs (...);
create table telemetry (...);
        """,
    )

    add_section_intro(
        doc,
        "8. API, Eventing, and Real-Time Communication",
        "The interaction model is hybrid. REST handles explicit user and agent requests. WebSocket handles live alerts. "
        "The event bus decouples producers from subscribers and keeps monitoring paths asynchronous.",
    )
    add_bullets(
        doc,
        [
            "REST endpoints: health, command, telemetry, memory, and events",
            "WebSocket endpoint: live alerts",
            "Event bus subscribers: alerts, future sync, future analytics, future long-running execution traces",
            "All async paths require explicit timeouts and graceful degradation",
        ],
    )

    add_section_intro(
        doc,
        "9. Capacity and Scalability",
        "The current workspace is intentionally sized for local or small-team deployment, but the architecture is designed to scale by replacing storage and messaging layers without rewriting the brains.",
    )
    add_kv_table(
        doc,
        [
            (
                "Local Baseline",
                "SQLite, in-process event bus, one backend instance, one or more edge agents",
            ),
            (
                "Team Scale",
                "Postgres, Redis, centralized websocket fan-out, per-device session inventory",
            ),
            (
                "Cloud Scale",
                "Managed Postgres, Redis or message bus, stateless API layer, separate telemetry ingestion workers",
            ),
            (
                "ML Scale",
                "External model inference service, batched anomaly training, feature store, retraining pipeline",
            ),
            (
                "Honeypot Scale",
                "Dedicated deception collectors, isolated network segments, normalized threat telemetry bus",
            ),
        ],
    )

    add_section_intro(
        doc,
        "10. Reliability and No Failure Rules",
        "JARVIS-X and DISHA must follow explicit no-failure constraints so the system fails visibly and safely instead of silently.",
    )
    add_rule_panel(
        doc,
        [
            "Every module must have a health check.",
            "Every decision must be logged.",
            "Every alert must have a reason.",
            "Every async flow must have a timeout.",
            "Every feature must degrade gracefully.",
        ],
    )
    add_bullets(
        doc,
        [
            "If ML anomaly detection is unavailable, the system falls back to statistical deviation analysis.",
            "If alert subscribers fail, the request still completes and the degraded path is logged.",
            "If a risky action is requested without confirmation, the system returns ASK instead of attempting execution.",
        ],
    )

    add_section_intro(
        doc,
        "11. Deployment Architecture",
        "The repo supports multiple maturity levels. DISHA web already has Docker Compose assets. JARVIS-X adds its own backend and agent containers. "
        "The target cloud architecture is a secure API control plane with stateless services, durable storage, and isolated telemetry ingestion.",
    )
    add_bullets(
        doc,
        [
            "Local run: `uvicorn` backend plus desktop agent plus web control plane as needed",
            "Container run: backend and agent through `disha/docker-compose.yml`",
            "Cloud-ready shape: API gateway, auth layer, message bus, data services, telemetry ingestion, alert channel",
            "Transport requirements: TLS, token auth, refresh rotation, secure sync keys, explicit device trust inventory",
        ],
    )

    add_section_intro(
        doc,
        "12. UI and Operator Experience",
        "The operator experience spans web, CLI, and mobile. Each surface must expose the same trust model and decision clarity, even if the interaction patterns differ.",
    )
    add_kv_table(
        doc,
        [
            (
                "Web Console",
                "Primary operational control plane with policy-aware workflows and audit visibility",
            ),
            ("CLI", "Trusted execution path for development and automation tasks"),
            (
                "Mobile",
                "Companion interface for chat, alerts, dashboard metrics, and settings",
            ),
            (
                "Voice",
                "Optional intent input layer routed through the same policy and planning system",
            ),
        ],
    )

    add_section_intro(
        doc,
        "13. Technical Decisions and Trade-Offs",
        "The current implementation and the target architecture intentionally favor clear, swappable boundaries over premature distribution.",
    )
    add_bullets(
        doc,
        [
            "SQLite is acceptable for baseline local persistence, but Postgres is the durable multi-user target.",
            "An in-process event bus is simple and fast for local use, but Redis streams, NATS, or Kafka are the next step for distributed deployments.",
            "The mobile app is a real scaffold, not a fully production-authenticated shipped client.",
            "Legacy Python modules are retained as intelligence assets until the target architecture stabilizes enough to absorb or retire them safely.",
        ],
    )

    add_section_intro(
        doc,
        "14. Limitations and Next Moves",
        "This document describes a complete architecture, but the repo is still in convergence. Not every legacy service has been migrated yet, and not every mobile or deception component is fully wired.",
    )
    add_numbered(
        doc,
        [
            "Converge duplicate Python service surfaces into explicit `brains/`, `services/`, and `edge/` ownership boundaries.",
            "Add automated tests for command routing, anomaly scoring, risk decisions, and alert fan-out.",
            "Promote SQLite-backed JARVIS-X persistence to Postgres when multi-user or cloud concurrency is required.",
            "Add isolated honeypot collectors and normalized deception event ingestion as a dedicated edge tier.",
            "Wire mobile screens to the live REST and WebSocket endpoints with real auth and sync management.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("Appendix A. Architecture Diagrams", level=1)
    add_code_block(
        doc,
        "Unified layered architecture (Mermaid)",
        """
flowchart TD
    Experience[Experience Layer]
    Orchestration[Orchestration Layer]
    Brains[Cognitive Brain Layer]
    Security[Security and Deception Layer]
    Telemetry[Intelligence and Telemetry Layer]
    Data[Data and Memory Layer]
    Edge[Edge and Sensor Layer]

    Experience --> Orchestration
    Orchestration --> Brains
    Brains --> Security
    Security --> Telemetry
    Telemetry --> Data
    Edge --> Telemetry
        """,
    )
    add_code_block(
        doc,
        "Threat and decision pipeline (Mermaid)",
        """
flowchart LR
    T[Telemetry or Honeypot Event]
    A[Anomaly Detection]
    R[Risk Engine]
    P[Policy Gate]
    D[Decision Engine]
    O1[Monitor]
    O2[Limit]
    O3[Isolate]
    O4[Ask]

    T --> A --> R --> P --> D
    D --> O1
    D --> O2
    D --> O3
    D --> O4
        """,
    )

    doc.sections[-1].start_type = WD_SECTION.NEW_PAGE
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
