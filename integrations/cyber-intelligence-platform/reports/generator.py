from docx import Document 

def generate_report(data):
    print("Generating report...")

    doc = Document()
    doc.add_heading("Cyber Intelligence Report", 0)

    for item in data:
        doc.add_paragraph(str(item))

    doc.save("report.docx")
    print("Report saved")
